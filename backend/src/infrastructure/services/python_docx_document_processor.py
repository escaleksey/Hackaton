import re
from copy import deepcopy
from dataclasses import replace
from io import BytesIO
from math import ceil

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.run import Run as DocxRun

from src.application.services.document_processor import (
    ContractDocumentProcessor,
    DownloadableContractFile,
    ParsedContractDocument,
)
from src.domain.entities.contract import ContractDraft
from src.domain.entities.document import (
    DEFAULT_PAGE_HEIGHT_PT,
    DEFAULT_PAGE_WIDTH_PT,
    DocumentLayout,
    DocumentPage,
    ParagraphBlock,
    TextRun,
    flatten_document_pages,
    normalize_document_pages,
)
from src.domain.entities.issue import ContractIssue

ALIGNMENT_TO_NAME = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

NAME_TO_ALIGNMENT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

SEVERITY_TO_HIGHLIGHT = {
    "high": "#ffe08a",
    "medium": "#a9e6b3",
    "low": "#9ad6ff",
}

DOCX_HIGHLIGHT_BY_COLOR = {
    "#ffe08a": WD_COLOR_INDEX.YELLOW,
    "#a9e6b3": WD_COLOR_INDEX.BRIGHT_GREEN,
    "#9ad6ff": WD_COLOR_INDEX.TURQUOISE,
}


class PythonDocxDocumentProcessor(ContractDocumentProcessor):
    def parse(
        self,
        filename: str,
        *,
        file_bytes: bytes | None = None,
        text: str | None = None,
    ) -> ParsedContractDocument:
        normalized_filename = filename.strip() or "contract.txt"
        lower_filename = normalized_filename.lower()

        if file_bytes is not None and lower_filename.endswith(".docx"):
            return self._parse_docx(normalized_filename, file_bytes)

        if file_bytes is not None:
            decoded_text = self._decode_text_file(file_bytes)
            return self._parse_text(normalized_filename, decoded_text)

        if text is not None:
            return self._parse_text(normalized_filename, text)

        raise ValueError("Передайте содержимое файла или текст договора.")

    def build_download(self, contract: ContractDraft) -> DownloadableContractFile:
        if contract.source_format == "docx":
            if contract.annotated_file_bytes is not None:
                filename = contract.filename
                if not filename.lower().endswith(".docx"):
                    filename = f"{filename}.docx"
                filename = f"{filename[:-5]}_corrected.docx"
                return DownloadableContractFile(
                    filename=filename,
                    content=contract.annotated_file_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            return self._build_docx(contract)

        filename = contract.filename
        if not filename.lower().endswith(".txt"):
            filename = f"{filename}.txt"
        filename = f"{filename[:-4]}_corrected.txt"

        return DownloadableContractFile(
            filename=filename,
            content=contract.current_text.encode("utf-8"),
            media_type="text/plain",
        )

    def build_annotated_source_download(
        self,
        *,
        filename: str,
        source_format: str,
        file_bytes: bytes | None,
        issues: list[ContractIssue],
    ) -> DownloadableContractFile | None:
        if source_format != "docx" or file_bytes is None:
            return None

        content = self._annotate_source_docx(file_bytes, issues)
        output_filename = filename if filename.lower().endswith(".docx") else f"{filename}.docx"
        output_filename = f"{output_filename[:-5]}_corrected.docx"
        return DownloadableContractFile(
            filename=output_filename,
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def annotate_pages(
        self,
        pages: list[DocumentPage],
        issues: list[ContractIssue],
    ) -> list[DocumentPage]:
        normalized_pages = normalize_document_pages(pages)
        issues_by_paragraph: dict[int, list[ContractIssue]] = {}

        for issue in issues:
            issues_by_paragraph.setdefault(issue.paragraph_index, []).append(issue)

        annotated_pages: list[DocumentPage] = []
        paragraph_index = 0

        for page in normalized_pages:
            annotated_blocks: list[ParagraphBlock] = []
            for block in page.blocks:
                paragraph_index += 1
                paragraph_issues = issues_by_paragraph.get(paragraph_index, [])
                annotated_blocks.append(
                    self._annotate_block(block, paragraph_issues) if paragraph_issues else block
                )

            annotated_pages.append(DocumentPage(number=page.number, blocks=annotated_blocks))

        return annotated_pages

    def _parse_text(self, filename: str, text: str) -> ParsedContractDocument:
        if not text.strip():
            raise ValueError("Contract text must not be empty.")

        blocks = self._text_to_blocks(text)
        layout = DocumentLayout()
        pages = self._paginate_blocks(blocks, layout)

        return ParsedContractDocument(
            filename=filename if filename.lower().endswith(".txt") else f"{filename}.txt",
            source_format="txt",
            text=text,
            pages=pages,
            document_layout=layout,
        )

    def _parse_docx(self, filename: str, file_bytes: bytes) -> ParsedContractDocument:
        try:
            document = Document(BytesIO(file_bytes))
        except Exception as error:
            raise ValueError("Не удалось прочитать DOCX-файл") from error

        layout = self._extract_layout(document)
        blocks: list[ParagraphBlock] = []
        split_before_indices: set[int] = set()
        split_after_indices: set[int] = set()

        for index, paragraph in enumerate(document.paragraphs):
            if paragraph.paragraph_format.page_break_before:
                split_before_indices.add(index)

            blocks.append(self._paragraph_to_block(paragraph, index))

            if self._paragraph_has_page_break(paragraph):
                split_after_indices.add(index)

        if not blocks:
            blocks = [ParagraphBlock(id="page-1-block-1", runs=[TextRun(text="")])]

        if split_before_indices or split_after_indices:
            pages = self._pages_from_explicit_breaks(
                blocks,
                split_before_indices,
                split_after_indices,
            )
        else:
            pages = self._paginate_blocks(blocks, layout)

        return ParsedContractDocument(
            filename=filename,
            source_format="docx",
            text=flatten_document_pages(pages),
            pages=pages,
            document_layout=layout,
            source_file_bytes=file_bytes,
        )

    def _build_docx(self, contract: ContractDraft) -> DownloadableContractFile:
        document = Document()
        self._apply_layout(document, contract.document_layout)
        pages = normalize_document_pages(contract.current_pages)

        first_paragraph = True
        for page_index, page in enumerate(pages):
            if page_index > 0:
                document.add_page_break()

            for block in page.blocks:
                paragraph = self._get_next_paragraph(document, first_paragraph)
                first_paragraph = False
                self._apply_paragraph_format(paragraph, block)

                if not block.runs:
                    paragraph.add_run("")
                    continue

                for run_data in block.runs:
                    run = paragraph.add_run(run_data.text)
                    run.bold = run_data.bold
                    run.italic = run_data.italic
                    run.underline = run_data.underline

                    if run_data.font_name:
                        run.font.name = run_data.font_name
                    if run_data.font_size_pt:
                        run.font.size = Pt(run_data.font_size_pt)
                    if run_data.color:
                        run.font.color.rgb = RGBColor.from_string(run_data.color.lstrip("#"))
                    if run_data.highlight_color:
                        run.font.highlight_color = DOCX_HIGHLIGHT_BY_COLOR.get(
                            run_data.highlight_color.lower(),
                            WD_COLOR_INDEX.YELLOW,
                        )

        buffer = BytesIO()
        document.save(buffer)
        filename = contract.filename
        if not filename.lower().endswith(".docx"):
            filename = f"{filename}.docx"
        filename = f"{filename[:-5]}_corrected.docx"

        return DownloadableContractFile(
            filename=filename,
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _decode_text_file(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8", "cp1251"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        raise ValueError("Не удалось определить кодировку файла. Используйте UTF-8 или CP1251")

    def _annotate_source_docx(self, file_bytes: bytes, issues: list[ContractIssue]) -> bytes:
        document = Document(BytesIO(file_bytes))
        issues_by_paragraph: dict[int, list[ContractIssue]] = {}

        for issue in issues:
            issues_by_paragraph.setdefault(issue.paragraph_index, []).append(issue)

        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            paragraph_issues = issues_by_paragraph.get(paragraph_index, [])
            if paragraph_issues:
                self._highlight_docx_paragraph(paragraph, paragraph_issues)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _text_to_blocks(self, text: str) -> list[ParagraphBlock]:
        lines = text.splitlines() or [text]
        blocks: list[ParagraphBlock] = []

        for index, line in enumerate(lines, start=1):
            blocks.append(
                ParagraphBlock(
                    id=f"page-1-block-{index}",
                    runs=[TextRun(text=line)],
                    alignment="left",
                )
            )

        return blocks

    def _extract_layout(self, document: DocxDocument) -> DocumentLayout:
        if not document.sections:
            return DocumentLayout()

        section = document.sections[0]
        return DocumentLayout(
            page_width_pt=self._length_to_pt(section.page_width) or DEFAULT_PAGE_WIDTH_PT,
            page_height_pt=self._length_to_pt(section.page_height) or DEFAULT_PAGE_HEIGHT_PT,
            margin_top_pt=self._length_to_pt(section.top_margin),
            margin_right_pt=self._length_to_pt(section.right_margin),
            margin_bottom_pt=self._length_to_pt(section.bottom_margin),
            margin_left_pt=self._length_to_pt(section.left_margin),
        )

    def _paragraph_to_block(self, paragraph, index: int) -> ParagraphBlock:
        style_font = self._get_style_font(paragraph.style)
        style_paragraph_format = self._get_style_paragraph_format(paragraph.style)
        runs: list[TextRun] = []
        for run in paragraph.runs:
            text = run.text or ""
            if text or not runs:
                color = self._resolve_run_color(run, style_font)
                runs.append(
                    TextRun(
                        text=text,
                        bold=self._resolve_run_flag(run.bold, style_font.bold),
                        italic=self._resolve_run_flag(run.italic, style_font.italic),
                        underline=self._resolve_run_flag(run.underline, style_font.underline),
                        font_name=run.font.name or style_font.name,
                        font_size_pt=self._length_to_pt(run.font.size or style_font.size),
                        color=f"#{color}" if color else None,
                        highlight_color=None,
                    )
                )

        if not runs:
            resolved_style_color = self._resolve_font_color(style_font)
            runs = [
                TextRun(
                    text="",
                    bold=self._resolve_run_flag(None, style_font.bold),
                    italic=self._resolve_run_flag(None, style_font.italic),
                    underline=self._resolve_run_flag(None, style_font.underline),
                    font_name=style_font.name,
                    font_size_pt=self._length_to_pt(style_font.size),
                    color=f"#{resolved_style_color}" if resolved_style_color else None,
                    highlight_color=None,
                )
            ]

        line_spacing_value = (
            paragraph.paragraph_format.line_spacing or style_paragraph_format.line_spacing
        )
        if hasattr(line_spacing_value, "pt"):
            line_spacing = line_spacing_value.pt
        elif isinstance(line_spacing_value, (int, float)):
            line_spacing = float(line_spacing_value)
        else:
            line_spacing = None

        return ParagraphBlock(
            id=f"page-1-block-{index + 1}",
            runs=runs,
            alignment=ALIGNMENT_TO_NAME.get(
                paragraph.alignment or style_paragraph_format.alignment,
                "left",
            ),
            style_name=paragraph.style.name if paragraph.style is not None else None,
            space_before_pt=self._length_to_pt(
                paragraph.paragraph_format.space_before or style_paragraph_format.space_before
            ),
            space_after_pt=self._length_to_pt(
                paragraph.paragraph_format.space_after or style_paragraph_format.space_after
            ),
            line_spacing=line_spacing,
        )

    def _paragraph_has_page_break(self, paragraph) -> bool:
        for run in paragraph.runs:
            for element in run._r.findall(qn("w:br")):
                if element.get(qn("w:type")) == "page":
                    return True
        return False

    def _pages_from_explicit_breaks(
        self,
        blocks: list[ParagraphBlock],
        split_before_indices: set[int],
        split_after_indices: set[int],
    ) -> list[DocumentPage]:
        pages: list[DocumentPage] = []
        current_blocks: list[ParagraphBlock] = []

        for index, block in enumerate(blocks):
            if index in split_before_indices and current_blocks:
                pages.append(DocumentPage(number=len(pages) + 1, blocks=current_blocks))
                current_blocks = []

            current_blocks.append(block)

            if index in split_after_indices:
                pages.append(DocumentPage(number=len(pages) + 1, blocks=current_blocks))
                current_blocks = []

        if current_blocks or not pages:
            pages.append(DocumentPage(number=len(pages) + 1, blocks=current_blocks))

        return normalize_document_pages(pages)

    def _paginate_blocks(
        self,
        blocks: list[ParagraphBlock],
        layout: DocumentLayout | None,
    ) -> list[DocumentPage]:
        target_layout = layout or DocumentLayout()
        available_height, content_width = self._get_available_body_size(target_layout)

        pages: list[DocumentPage] = []
        current_blocks: list[ParagraphBlock] = []
        current_height = 0.0

        for block in blocks:
            block_height = self._estimate_block_height(block, content_width)
            if current_blocks and current_height + block_height > available_height:
                pages.append(DocumentPage(number=len(pages) + 1, blocks=current_blocks))
                current_blocks = []
                current_height = 0.0

            current_blocks.append(block)
            current_height += block_height

        if current_blocks or not pages:
            pages.append(DocumentPage(number=len(pages) + 1, blocks=current_blocks))

        return normalize_document_pages(pages)

    def _estimate_block_height(self, block: ParagraphBlock, content_width_pt: float) -> float:
        sample_run = block.runs[0] if block.runs else TextRun(text="")
        font_size = sample_run.font_size_pt or 12.0
        line_height = max(font_size * 1.42, 16.0)
        average_char_width = max(font_size * 0.68, 6.8)
        chars_per_line = max(24, int(content_width_pt / average_char_width))
        text_lines = (block.text or " ").splitlines() or [" "]
        estimated_lines = sum(max(1, ceil(len(line) / chars_per_line)) for line in text_lines)
        spacing_before = block.space_before_pt or 0.0
        spacing_after = block.space_after_pt or 6.0
        explicit_line_spacing = block.line_spacing or line_height
        if explicit_line_spacing <= 5:
            explicit_line_spacing = line_height * explicit_line_spacing
        if block.style_name and "heading" in block.style_name.lower():
            explicit_line_spacing += 4.0
            spacing_after += 8.0
        return spacing_before + spacing_after + estimated_lines * explicit_line_spacing

    def _apply_layout(self, document: DocxDocument, layout: DocumentLayout | None) -> None:
        if layout is None or not document.sections:
            return

        section = document.sections[0]
        section.page_width = Pt(layout.page_width_pt)
        section.page_height = Pt(layout.page_height_pt)

        if layout.margin_top_pt is not None:
            section.top_margin = Pt(layout.margin_top_pt)
        if layout.margin_right_pt is not None:
            section.right_margin = Pt(layout.margin_right_pt)
        if layout.margin_bottom_pt is not None:
            section.bottom_margin = Pt(layout.margin_bottom_pt)
        if layout.margin_left_pt is not None:
            section.left_margin = Pt(layout.margin_left_pt)

    def _get_next_paragraph(self, document: DocxDocument, first_paragraph: bool):
        if first_paragraph and len(document.paragraphs) == 1 and not document.paragraphs[0].text:
            paragraph = document.paragraphs[0]
            paragraph.clear()
            return paragraph
        return document.add_paragraph()

    def _apply_paragraph_format(self, paragraph, block: ParagraphBlock) -> None:
        paragraph.alignment = NAME_TO_ALIGNMENT.get(block.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        if block.style_name:
            try:
                paragraph.style = block.style_name
            except KeyError:
                pass

        if block.space_before_pt is not None:
            paragraph.paragraph_format.space_before = Pt(block.space_before_pt)
        if block.space_after_pt is not None:
            paragraph.paragraph_format.space_after = Pt(block.space_after_pt)
        if block.line_spacing is not None:
            if block.line_spacing <= 5:
                paragraph.paragraph_format.line_spacing = block.line_spacing
            else:
                paragraph.paragraph_format.line_spacing = Pt(block.line_spacing)

    def _length_to_pt(self, value) -> float | None:
        return value.pt if value is not None else None

    def _get_available_body_size(self, layout: DocumentLayout) -> tuple[float, float]:
        page_height = layout.page_height_pt or DEFAULT_PAGE_HEIGHT_PT
        page_width = layout.page_width_pt or DEFAULT_PAGE_WIDTH_PT
        vertical_margins = (layout.margin_top_pt or 56.0) + (layout.margin_bottom_pt or 56.0)
        horizontal_margins = (layout.margin_left_pt or 56.0) + (layout.margin_right_pt or 56.0)

        # Reserve extra space for headers/footers and reduce over-merging of pages.
        available_height = max(360.0, page_height - vertical_margins - 72.0)
        content_width = max(220.0, page_width - horizontal_margins - 24.0)
        return available_height, content_width

    def _get_style_font(self, style):
        current_style = style
        while current_style is not None:
            font = current_style.font
            if (
                font is not None
                and (
                    font.name is not None
                    or font.size is not None
                    or font.bold is not None
                    or font.italic is not None
                    or font.underline is not None
                    or font.color.rgb is not None
                )
            ):
                return font
            current_style = current_style.base_style

        return style.font if style is not None else Document().styles["Normal"].font

    def _get_style_paragraph_format(self, style):
        current_style = style
        while current_style is not None:
            paragraph_format = current_style.paragraph_format
            if (
                paragraph_format is not None
                and (
                    paragraph_format.alignment is not None
                    or paragraph_format.space_before is not None
                    or paragraph_format.space_after is not None
                    or paragraph_format.line_spacing is not None
                )
            ):
                return paragraph_format
            current_style = current_style.base_style

        if style is not None:
            return style.paragraph_format
        return Document().styles["Normal"].paragraph_format

    def _resolve_run_flag(self, run_value, style_value) -> bool:
        if run_value is not None:
            return bool(run_value)
        if style_value is not None:
            return bool(style_value)
        return False

    def _resolve_font_color(self, font) -> str | None:
        if font is None or font.color is None or font.color.rgb is None:
            return None
        return str(font.color.rgb)

    def _resolve_run_color(self, run, style_font) -> str | None:
        if run.font.color is not None and run.font.color.rgb is not None:
            return str(run.font.color.rgb)
        return self._resolve_font_color(style_font)

    def _annotate_block(
        self,
        block: ParagraphBlock,
        issues: list[ContractIssue],
    ) -> ParagraphBlock:
        ranges = self._resolve_highlight_ranges(block.text, issues)
        if not ranges:
            return block

        boundaries = {0, len(block.text)}
        for start, end, _ in ranges:
            boundaries.add(start)
            boundaries.add(end)
        sorted_boundaries = sorted(boundaries)

        current_position = 0
        annotated_runs: list[TextRun] = []

        for run in block.runs:
            run_text = run.text or ""
            run_start = current_position
            run_end = current_position + len(run_text)
            current_position = run_end

            if run_start == run_end:
                annotated_runs.append(run)
                continue

            run_boundaries = [run_start]
            run_boundaries.extend(
                boundary for boundary in sorted_boundaries if run_start < boundary < run_end
            )
            run_boundaries.append(run_end)

            for start, end in zip(run_boundaries, run_boundaries[1:], strict=False):
                piece = run_text[start - run_start : end - run_start]
                if not piece:
                    continue
                annotated_runs.append(
                    replace(
                        run,
                        text=piece,
                        highlight_color=self._resolve_highlight_color(start, end, ranges),
                    )
                )

        if annotated_runs:
            return replace(block, runs=annotated_runs)
        return replace(
            block,
            runs=[replace(block.runs[0], highlight_color="#ffe08a")],
        )

    def _resolve_highlight_ranges(
        self,
        text: str,
        issues: list[ContractIssue],
    ) -> list[tuple[int, int, str]]:
        ranges: list[tuple[int, int, str]] = []

        for issue in sorted(issues, key=lambda item: self._severity_rank(item.severity)):
            fragment = (issue.fragment or "").strip()
            color = SEVERITY_TO_HIGHLIGHT.get(
                issue.severity.lower(),
                SEVERITY_TO_HIGHLIGHT["medium"],
            )

            if not fragment:
                if text:
                    ranges.append((0, len(text), color))
                continue

            matches = self._find_fragment_ranges(text, fragment)
            if matches:
                ranges.extend((start, end, color) for start, end in matches)
            elif text:
                ranges.append((0, len(text), color))

        return ranges

    def _find_fragment_ranges(self, text: str, fragment: str) -> list[tuple[int, int]]:
        if not text or not fragment:
            return []

        lowered_text = text.lower()
        lowered_fragment = fragment.lower()
        matches: list[tuple[int, int]] = []
        cursor = 0

        while True:
            index = lowered_text.find(lowered_fragment, cursor)
            if index < 0:
                break
            matches.append((index, index + len(fragment)))
            cursor = index + len(fragment)

        if matches:
            return matches

        fragment_parts = [part for part in fragment.split() if part]
        if not fragment_parts:
            return []

        pattern = r"\s+".join(re.escape(part) for part in fragment_parts)
        return [(match.start(), match.end()) for match in re.finditer(pattern, text, re.IGNORECASE)]

    def _resolve_highlight_color(
        self,
        start: int,
        end: int,
        ranges: list[tuple[int, int, str]],
    ) -> str | None:
        for range_start, range_end, color in ranges:
            if start < range_end and end > range_start:
                return color
        return None

    def _severity_rank(self, severity: str) -> int:
        order = {"high": 0, "medium": 1, "low": 2}
        return order.get(severity.lower(), 99)

    def _highlight_docx_paragraph(self, paragraph, issues: list[ContractIssue]) -> None:
        ranges = self._resolve_highlight_ranges(paragraph.text, issues)
        if not ranges:
            return

        boundaries = {0, len(paragraph.text)}
        for start, end, _ in ranges:
            boundaries.add(start)
            boundaries.add(end)
        sorted_boundaries = sorted(boundaries)

        current_position = 0
        original_runs = list(paragraph.runs)

        for run in original_runs:
            run_text = run.text or ""
            run_start = current_position
            run_end = current_position + len(run_text)
            current_position = run_end

            if not run_text:
                continue

            run_boundaries = [run_start]
            run_boundaries.extend(
                boundary for boundary in sorted_boundaries if run_start < boundary < run_end
            )
            run_boundaries.append(run_end)

            segments: list[tuple[str, WD_COLOR_INDEX | None]] = []
            for start, end in zip(run_boundaries, run_boundaries[1:], strict=False):
                piece = run_text[start - run_start : end - run_start]
                if piece:
                    segments.append(
                        (
                            piece,
                            self._resolve_docx_highlight_color(start, end, ranges, run),
                        )
                    )

            if not segments:
                continue

            run.text = segments[0][0]
            run.font.highlight_color = segments[0][1]

            insert_after = run
            for piece, highlight_color in segments[1:]:
                insert_after = self._clone_run_after(insert_after, piece, highlight_color)

    def _clone_run_after(
        self,
        run,
        text: str,
        highlight_color: WD_COLOR_INDEX | None,
    ):
        cloned_xml = deepcopy(run._r)
        run._r.addnext(cloned_xml)
        cloned_run = DocxRun(cloned_xml, run._parent)
        cloned_run.text = text
        cloned_run.font.highlight_color = highlight_color
        return cloned_run

    def _resolve_docx_highlight_color(
        self,
        start: int,
        end: int,
        ranges: list[tuple[int, int, str]],
        run,
    ) -> WD_COLOR_INDEX | None:
        for range_start, range_end, color in ranges:
            if start < range_end and end > range_start:
                return DOCX_HIGHLIGHT_BY_COLOR.get(color.lower(), WD_COLOR_INDEX.YELLOW)
        return run.font.highlight_color
