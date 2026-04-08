from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from fastapi.testclient import TestClient

from src.application.services.legal_ai_pipeline import (
    LegalAiPipeline,
    SemanticAnalysisResult,
)
from src.domain.entities.analysis_warning import ContractAnalysisWarning
from src.application.use_cases.create_contract_draft import CreateContractDraftUseCase
from src.domain.entities.issue import ContractIssue
from src.main import app
from src.presentation.api.dependencies import (
    get_contract_repository,
    get_create_contract_draft_use_case,
    get_document_processor,
)


class StubLegalAiPipeline(LegalAiPipeline):
    def __init__(
        self,
        issues: list[ContractIssue],
        warnings: list[ContractAnalysisWarning] | None = None,
    ) -> None:
        self._issues = issues
        self._warnings = warnings or []

    def analyze(self, document) -> SemanticAnalysisResult:
        return SemanticAnalysisResult(issues=self._issues, warnings=self._warnings)


@pytest.fixture(autouse=True)
def clear_contract_repository() -> None:
    get_contract_repository()._storage.clear()
    app.dependency_overrides.clear()


@pytest.fixture
def client() -> TestClient:
    return TestClient(app)


def build_docx_file() -> bytes:
    document = Document()
    document.add_heading("Заголовок договора", level=1)

    paragraph = document.add_paragraph()
    run = paragraph.add_run("Договор действует с ")
    run.italic = True
    run = paragraph.add_run("25.03.2023")
    run.bold = True
    paragraph.add_run(" по ")
    run = paragraph.add_run("02.02.2023")
    run.underline = True
    paragraph.add_run(".")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def override_create_use_case(
    issues: list[ContractIssue],
    warnings: list[ContractAnalysisWarning] | None = None,
):
    def _factory() -> CreateContractDraftUseCase:
        return CreateContractDraftUseCase(
            get_contract_repository(),
            get_document_processor(),
            StubLegalAiPipeline(issues, warnings),
        )

    return _factory


def test_upload_apply_and_download_text_contract(client: TestClient) -> None:
    app.dependency_overrides[get_create_contract_draft_use_case] = override_create_use_case([])

    upload_response = client.post(
        "/api/v1/contracts/upload",
        data={"text": "Исходный договор"},
    )

    assert upload_response.status_code == 201
    draft = upload_response.json()
    assert draft["source_format"] == "txt"
    assert draft["issues"] == []

    apply_response = client.post(
        f"/api/v1/contracts/{draft['id']}/apply",
        json={"corrected_text": "Исправленный договор"},
    )

    assert apply_response.status_code == 200
    assert apply_response.json()["corrected_text"] == "Исправленный договор"

    download_response = client.get(f"/api/v1/contracts/{draft['id']}/download")
    assert download_response.status_code == 200
    assert download_response.text == "Исправленный договор"
    assert "attachment" in download_response.headers["content-disposition"]


def test_upload_docx_returns_issues_and_highlighted_preview(client: TestClient) -> None:
    app.dependency_overrides[get_create_contract_draft_use_case] = override_create_use_case(
        [
            ContractIssue(
                paragraph_index=2,
                fragment="с 25.03.2023 по 02.02.2023",
                type="DATE_CONFLICT",
                severity="high",
                confidence="high",
                explanation="Дата окончания раньше даты начала.",
                suggestion="Уточнить дату окончания договора.",
            )
        ]
    )

    upload_response = client.post(
        "/api/v1/contracts/upload",
        files={
            "file": (
                "contract.docx",
                build_docx_file(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201
    draft = upload_response.json()
    assert draft["source_format"] == "docx"
    assert draft["issues"][0]["type"] == "DATE_CONFLICT"
    assert any(
        run["highlight_color"]
        for run in draft["corrected_pages"][0]["blocks"][1]["runs"]
        if "25.03.2023" in run["text"] or "02.02.2023" in run["text"] or "по" in run["text"]
    )


def test_upload_returns_analysis_warning_when_llm_falls_back(client: TestClient) -> None:
    app.dependency_overrides[get_create_contract_draft_use_case] = override_create_use_case(
        [],
        [
            ContractAnalysisWarning(
                code="llm_insufficient_quota",
                message="Показаны только локальные rule-based проверки.",
            )
        ],
    )

    upload_response = client.post(
        "/api/v1/contracts/upload",
        data={"text": "Текст договора"},
    )

    assert upload_response.status_code == 201
    draft = upload_response.json()
    assert draft["warnings"] == [
        {
            "code": "llm_insufficient_quota",
            "message": "Показаны только локальные rule-based проверки.",
        }
    ]


def test_download_marked_docx_preserves_original_styles(client: TestClient) -> None:
    app.dependency_overrides[get_create_contract_draft_use_case] = override_create_use_case(
        [
            ContractIssue(
                paragraph_index=2,
                fragment="с 25.03.2023 по 02.02.2023",
                type="DATE_CONFLICT",
                severity="high",
                confidence="high",
                explanation="Дата окончания раньше даты начала.",
                suggestion="Уточнить дату окончания договора.",
            )
        ]
    )

    upload_response = client.post(
        "/api/v1/contracts/upload",
        files={
            "file": (
                "contract.docx",
                build_docx_file(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201
    draft = upload_response.json()

    download_response = client.get(f"/api/v1/contracts/{draft['id']}/download")
    assert download_response.status_code == 200

    downloaded_document = Document(BytesIO(download_response.content))
    heading = downloaded_document.paragraphs[0]
    body = downloaded_document.paragraphs[1]

    assert heading.style.name == "Heading 1"
    assert any(run.italic for run in body.runs)
    assert any(run.bold for run in body.runs)
    assert any(run.underline for run in body.runs)
    assert any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in body.runs)


def test_apply_and_download_docx_preserves_heading_formatting_and_highlights(
    client: TestClient,
) -> None:
    app.dependency_overrides[get_create_contract_draft_use_case] = override_create_use_case(
        [
            ContractIssue(
                paragraph_index=2,
                fragment="с 25.03.2023 по 02.02.2023",
                type="DATE_CONFLICT",
                severity="high",
                confidence="high",
                explanation="Дата окончания раньше даты начала.",
                suggestion="Уточнить дату окончания договора.",
            )
        ]
    )

    upload_response = client.post(
        "/api/v1/contracts/upload",
        files={
            "file": (
                "contract.docx",
                build_docx_file(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201
    draft = upload_response.json()
    corrected_pages = draft["corrected_pages"]

    corrected_pages[0]["blocks"][0]["alignment"] = "center"
    corrected_pages[0]["blocks"][0]["style_name"] = "Heading 1"
    corrected_pages[0]["blocks"][0]["runs"] = [
        {
            "text": "Исправленный заголовок",
            "bold": True,
            "italic": False,
            "underline": False,
            "font_name": "Arial",
            "font_size_pt": 18,
            "color": None,
            "highlight_color": None,
        }
    ]

    apply_response = client.post(
        f"/api/v1/contracts/{draft['id']}/apply",
        json={"corrected_pages": corrected_pages},
    )

    assert apply_response.status_code == 200

    download_response = client.get(f"/api/v1/contracts/{draft['id']}/download")
    assert download_response.status_code == 200
    assert (
        download_response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    downloaded_document = Document(BytesIO(download_response.content))
    heading = downloaded_document.paragraphs[0]
    body = downloaded_document.paragraphs[1]

    assert heading.text == "Исправленный заголовок"
    assert heading.style.name == "Heading 1"
    assert heading.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert heading.runs[0].font.size is not None
    assert round(heading.runs[0].font.size.pt) == 18
    assert any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in body.runs)


def test_upload_returns_issue_replacement_field(client: TestClient) -> None:
    app.dependency_overrides[get_create_contract_draft_use_case] = override_create_use_case(
        [
            ContractIssue(
                paragraph_index=1,
                fragment="Заказчик",
                type="TERM_MISUSE",
                severity="medium",
                confidence="high",
                explanation="Термин стороны используется непоследовательно.",
                suggestion="Унифицировать термин по всему договору.",
                replacement="Клиент",
            )
        ]
    )

    upload_response = client.post(
        "/api/v1/contracts/upload",
        data={"text": "Клиент подписывает договор. Заказчик оплачивает услуги."},
    )

    assert upload_response.status_code == 201
    draft = upload_response.json()
    assert draft["issues"][0]["type"] == "TERM_MISUSE"
    assert draft["issues"][0]["replacement"] == "Клиент"
